"""
processing.py — Advanced Document Processing for CollegeAI
Features:
  - Multi-engine OCR (pytesseract + easyocr fallback)
  - Parallel page processing for PDFs
  - Image pre-processing pipeline (deskew, denoise, threshold)
  - Enhanced text extraction with layout awareness
  - Result parsing: table extraction first, then fixed-coordinate fallback (from working script)
  - BFS web scraper (unchanged)
  - PDFs saved without hash prefix (original filename only)
"""
import re
import os
import io
import json
import time
import hashlib
import requests
import pdfplumber
import numpy as np
import xml.etree.ElementTree as ET
import concurrent.futures

from PIL import Image, ImageFilter, ImageOps
from urllib.parse import urljoin, urlparse
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from sentence_transformers import SentenceTransformer
from collections import deque

from app.document.faiss_manager import get_index, save_index
from app.document.result_extractor import (
    detect_semester_from_text,
    parse_result_document,
    students_to_result_chunks,
)
from app.database import SessionLocal
from app.models.chunk import DocumentChunk
from app.logger import logger


# ─── OCR Engine Setup (unchanged) ────────────────────────────────────────────
OCR_AVAILABLE   = False
EASYOCR_AVAILABLE = False
_easyocr_reader = None

try:
    import pytesseract
    from pdf2image import convert_from_bytes

    TESSERACT_PATH = os.getenv("TESSERACT_PATH", r"C:\Program Files\Tesseract-OCR\tesseract.exe")
    POPPLER_PATH   = os.getenv("POPPLER_PATH",   r"C:\poppler-25.12.0\Library\bin")

    if TESSERACT_PATH and os.path.exists(TESSERACT_PATH):
        pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH

    OCR_AVAILABLE = True
    logger.info("pytesseract OCR available")
except ImportError:
    logger.warning("pytesseract / pdf2image not installed — primary OCR disabled")

try:
    import easyocr
    EASYOCR_AVAILABLE = True
    logger.info("easyocr available as fallback OCR engine")
except ImportError:
    pass


def _get_easyocr_reader():
    global _easyocr_reader
    if _easyocr_reader is None and EASYOCR_AVAILABLE:
        import easyocr
        _easyocr_reader = easyocr.Reader(["en"], gpu=False, verbose=False)
    return _easyocr_reader


# ─── Scraper Config (unchanged) ──────────────────────────────────────────────
SAME_DOMAIN_ONLY    = True
REQUEST_DELAY       = 0.3
REQUEST_TIMEOUT     = 20
MIN_IMAGE_SIZE_KB   = 10
MAX_PAGES_PER_CRAWL = int(os.getenv("MAX_PAGES_PER_CRAWL", 500))
MAX_CRAWL_DEPTH     = int(os.getenv("MAX_CRAWL_DEPTH", 1))

# Comma-separated domains or URL substrings to never visit.
# Example in .env:  SCRAPE_BLOCKED_DOMAINS=/university-syllabus,/qp-ds,facebook.com
# Read lazily on every check so changes take effect without restart.
def _get_blocked_domains() -> set:
    raw = os.getenv("SCRAPE_BLOCKED_DOMAINS", "")
    return {d.strip().lower() for d in raw.split(",") if d.strip()}
MAX_HTML_CHARS      = 500_000
MAX_DOC_CHARS       = 2_000_000

# Parallel OCR workers
OCR_WORKERS = int(os.getenv("OCR_WORKERS", 4))
PDF_WORKERS = int(os.getenv("PDF_WORKERS", 4))

ALLOWED_EXTENSIONS = {
    "pdf", "docx", "txt", "xml",
    "jpg", "jpeg", "png", "gif", "webp", "bmp"
}

SKIP_URL_PATTERNS = [
    r"\.(css|js|woff|woff2|ttf|eot|ico|svg|mp4|mp3|zip|rar)(\?|$)",
    r"(logout|login|wp-admin|wp-login|cart|checkout|feed|rss)",
    r"#",
    r"javascript:",
    r"mailto:",
    r"tel:",
]

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/120.0.0.0 Safari/537.36"
    ),
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "en-US,en;q=0.5",
}


# ─── Embedding Model (singleton) ──────────────────────────────────────────────
_model = None

def get_model() -> SentenceTransformer:
    global _model
    if _model is None:
        _model = SentenceTransformer("all-MiniLM-L6-v2")
    return _model


# ─── Classifiers (unchanged) ──────────────────────────────────────────────────
DEPARTMENT_PATTERNS = {
    "DS":    [r"\bdata[\s\-_]?science\b", r"\bCSE[\s\-_]?DS\b"],
    "CE":    [r"\bcomputer[\s\-_]?eng(ineering)?\b", r"\bCSE\b(?![\s\-_]DS)"],
    "MECH":  [r"\bmechanical\b", r"\bMECH\b"],
    "CIVIL": [r"\bcivil\b"],
    "EXTC":  [r"\bEXTC\b", r"\belectronics\b", r"\btelecommunication\b"],
    "IT":    [r"\binformation[\s\-_]?technology\b", r"\bIT\b"],
    "ME":    [r"\bmaster[\s\-_]?of[\s\-]?engineering\b"],
}

CONTENT_PATTERNS = {
    "RESULT":   [r"\bresult\b", r"\bgazette\b", r"\bmarksheet\b",
                 r"\bDS\s*\d{4}\b", r"\bpassed\b", r"\bfailed\b",
                 r"\bseat[\s\-]?no\b", r"\btotal[\s\-]?marks\b", r"\bSGPI\b",
                 r"c[\s\-]?scheme[\s\-]?reg"],
    "NOTICE":   [r"\bnotice\b", r"\bcircular\b", r"\bannouncement\b",
                 r"\bexam[\s\-]?form\b", r"\bKT\b", r"\btime[\s\-]?table\b",
                 r"\bschedule\b", r"\bhall[\s\-]?ticket\b"],
    "SYLLABUS": [r"\bsyllabus\b", r"\bcurriculum\b", r"\bcourse[\s\-]?outline\b",
                 r"\bcourse[\s\-]?structure\b", r"\bteaching[\s\-]?scheme\b",
                 r"\bexamination[\s\-]?scheme\b", r"\blearning[\s\-]?outcomes?\b",
                 r"\bcourse[\s\-]?outcomes?\b", r"\bunit\s*[-:]?\s*(?:\d+|[ivx]+)\b",
                 r"\bmodule\s*[-:]?\s*\d+\b", r"\bPO\b", r"\bPSO\b", r"\bCO\b"],
    "FACULTY":  [r"\bfaculty\b", r"\bprofessor\b", r"\bstaff\b",
                 r"\bHOD\b", r"\bdesignation\b", r"\bhead[\s\-]?of[\s\-]?department\b"],
    "FEE":      [r"\bfee[\s\-]?structure\b", r"\btuition\b", r"\bscholarship\b",
                 r"\bfee[\s\-]?detail\b"],
    "PLACEMENT":[r"\bplacement\b", r"\brecruit\b", r"\bcampus[\s\-]?drive\b"],
    "ADMISSION":[r"\badmission\b", r"\beligibility\b", r"\bapply\b"],
    "NIRF":     [r"\bNIRF\b", r"\branking\b", r"\baccreditation\b",
                 r"\bNAAC\b", r"\bNBA\b"],
}

SEMESTER_PATTERN = re.compile(
    r"SEM(?:ESTER)?[\s\-]*([IVX]+|\d+)|Semester[\s:]+([IVX]+|\d+)",
    re.IGNORECASE
)

STATUS_MAP = {
    "PASS": "PASS", "P": "PASS",
    "FAIL": "FAIL", "F": "FAIL",
    "PF":   "PASS WITH FAIL",
    "PASSWITHFAIL": "PASS WITH FAIL",
}


def classify_chunk(text: str, source_url: str = "") -> dict:
    combined  = (source_url + " " + text[:2000]).upper()
    url_lower = source_url.lower()

    if "faculty" in url_lower or "staff" in url_lower:
        content_type = "FACULTY"
    elif "result" in url_lower or "gazette" in url_lower or "c-scheme-reg" in url_lower:
        content_type = "RESULT"
    elif (
        "syllabus" in url_lower
        or "curriculum" in url_lower
        or "course-structure" in url_lower
        or "teaching-scheme" in url_lower
    ):
        content_type = "SYLLABUS"
    elif "notice" in url_lower or "circular" in url_lower or "announcement" in url_lower:
        content_type = "NOTICE"
    elif "fee" in url_lower or "tuition" in url_lower:
        content_type = "FEE"
    elif "placement" in url_lower or "recruit" in url_lower:
        content_type = "PLACEMENT"
    elif "admission" in url_lower:
        content_type = "ADMISSION"
    elif "nirf" in url_lower or "naac" in url_lower:
        content_type = "NIRF"
    else:
        content_type = "GENERAL"
        for ctype, patterns in CONTENT_PATTERNS.items():
            if any(re.search(p, combined, re.IGNORECASE) for p in patterns):
                content_type = ctype
                break

    department = "GENERAL"
    for dept, patterns in DEPARTMENT_PATTERNS.items():
        if any(re.search(p, combined, re.IGNORECASE) for p in patterns):
            department = dept
            break

    sem_match = SEMESTER_PATTERN.search(combined)
    semester  = None
    if sem_match:
        val      = sem_match.group(1) or sem_match.group(2)
        semester = "SEM-" + val.upper()

    return {"department": department, "content_type": content_type, "semester": semester}


# ─── Advanced Image Pre-processing (unchanged) ────────────────────────────────
def _preprocess_image(img: Image.Image) -> Image.Image:
    img = img.convert("L")
    w, h = img.size
    if w < 1500 or h < 1500:
        scale = max(1500 / w, 1500 / h, 1)
        img   = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
    import numpy as _np
    arr   = _np.array(img)
    hist, bins = _np.histogram(arr.flatten(), bins=256, range=(0, 256))
    total = arr.size
    sumB  = 0; wB = 0; wF = 0
    maxVar = 0; threshold = 128
    total_sum = _np.dot(_np.arange(256), hist)
    for i in range(256):
        wB += hist[i]
        if wB == 0:
            continue
        wF = total - wB
        if wF == 0:
            break
        sumB += i * hist[i]
        mB = sumB / wB
        mF = (total_sum - sumB) / wF
        var = wB * wF * (mB - mF) ** 2
        if var > maxVar:
            maxVar    = var
            threshold = i
    img = img.point(lambda p: 255 if p > threshold else 0)
    img = img.filter(ImageFilter.SHARPEN)
    return img


def _deskew_image(img: Image.Image) -> Image.Image:
    try:
        import numpy as _np
        arr = _np.array(img.convert("L"))
        thresh = (arr < 128).astype(_np.uint8)
        h, w   = thresh.shape
        angles = _np.linspace(-5, 5, 21)
        best_angle = 0
        best_score = -1
        for angle in angles:
            rotated = img.rotate(angle, expand=False, fillcolor=255)
            arr_r   = _np.array(rotated.convert("L"))
            score   = _np.var(_np.sum((arr_r < 128), axis=1))
            if score > best_score:
                best_score = score
                best_angle = angle
        if abs(best_angle) > 0.5:
            img = img.rotate(best_angle, expand=False, fillcolor=255)
    except Exception:
        pass
    return img


def _ocr_single_image(img: Image.Image, page_num: int = 0) -> str:
    img = _preprocess_image(img)
    img = _deskew_image(img)
    if OCR_AVAILABLE:
        try:
            config = ("--oem 3 --psm 6 -c tessedit_char_whitelist=ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789.,/:()-_ \n")
            text = pytesseract.image_to_string(img, lang="eng", config=config)
            if text and len(text.strip()) > 20:
                return text
            text = pytesseract.image_to_string(img, lang="eng")
            if text and len(text.strip()) > 10:
                return text
        except Exception as e:
            logger.warning(f"[OCR] pytesseract page {page_num}: {e}")
    reader = _get_easyocr_reader()
    if reader:
        try:
            import numpy as _np
            results = reader.readtext(_np.array(img), detail=0, paragraph=True)
            text    = "\n".join(results)
            if text.strip():
                return text
        except Exception as e:
            logger.warning(f"[OCR] easyocr page {page_num}: {e}")
    return ""


# ─── PDF Text Extraction (unchanged) ─────────────────────────────────────────
def _extract_page_text(args) -> tuple[int, str]:
    """
    Extract text from a single PDF page with three fallback tiers:
      1. Full extraction (text + tables) — normal path.
      2. Text-only (skip tables) — if tables blow the buffer.
      3. Cropped-region extraction — if the page itself is too large
         (common on scanned A3/poster pages embedded in PDFs).
    """
    page_idx, pdf_bytes = args
    try:
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            if page_idx >= len(pdf.pages):
                return page_idx, ""
            page = pdf.pages[page_idx]

            # ── Tier 1: full extraction ──────────────────────────────────
            try:
                text = page.extract_text() or ""
                try:
                    for table in page.extract_tables():
                        if not table:
                            continue
                        for row in table:
                            row_text = "\t".join(str(c or "").strip() for c in row)
                            if row_text.strip():
                                text += "\n" + row_text
                except Exception:
                    pass  # table extraction optional
                return page_idx, text
            except Exception as e1:
                if "output buffer" not in str(e1).lower() and "allocate" not in str(e1).lower():
                    raise  # unexpected error — re-raise so outer handler logs it

            # ── Tier 2: text-only (skip tables, avoid large image decode) ─
            logger.debug(f"[PDF] Page {page_idx}: buffer error on full extract, trying text-only")
            try:
                text = page.extract_text(x_tolerance=3, y_tolerance=3) or ""
                return page_idx, text
            except Exception as e2:
                if "output buffer" not in str(e2).lower() and "allocate" not in str(e2).lower():
                    raise

            # ── Tier 3: crop page into horizontal bands and extract each ──
            logger.debug(f"[PDF] Page {page_idx}: text-only failed, trying banded crop")
            try:
                w, h = float(page.width), float(page.height)
                bands = 4
                band_h = h / bands
                parts = []
                for b in range(bands):
                    y0 = b * band_h
                    y1 = min((b + 1) * band_h, h)
                    try:
                        cropped = page.crop((0, y0, w, y1))
                        part = cropped.extract_text(x_tolerance=3, y_tolerance=3) or ""
                        if part.strip():
                            parts.append(part)
                    except Exception:
                        pass  # skip unreadable band
                return page_idx, "\n".join(parts)
            except Exception:
                pass

    except Exception as e:
        logger.warning(f"[PDF] Page {page_idx} extraction error: {e}")

    return page_idx, ""


def _ocr_page_worker(args) -> tuple[int, str]:
    page_idx, img_bytes = args
    try:
        img  = Image.open(io.BytesIO(img_bytes)).convert("RGB")
        text = _ocr_single_image(img, page_num=page_idx)
        return page_idx, text
    except Exception as e:
        logger.warning(f"[OCR] Page {page_idx} worker error: {e}")
        return page_idx, ""


def _extract_pdf_bytes(data: bytes) -> str:
    """
    Extract text from all pages of a PDF.

    Worker count is scaled down for large PDFs (>40 pages) to avoid
    simultaneous buffer allocation errors -- each parallel worker keeps
    the entire PDF in memory while decoding its page.

    Pages that still fail in the parallel pass (buffer errors on very
    large/image-heavy pages) are retried serially so they get the full
    banded-crop fallback path in _extract_page_text.
    """
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            n_pages = len(pdf.pages)
        if n_pages == 0:
            return ""

        # Fewer threads = less peak RAM for large PDFs
        if n_pages > 80:
            workers = max(1, PDF_WORKERS - 2)
        elif n_pages > 40:
            workers = max(1, PDF_WORKERS - 1)
        else:
            workers = PDF_WORKERS

        page_args  = [(i, data) for i in range(n_pages)]
        page_texts = [""] * n_pages
        failed_pages: list[int] = []

        with concurrent.futures.ThreadPoolExecutor(max_workers=workers) as exe:
            futures = {exe.submit(_extract_page_text, arg): arg[0] for arg in page_args}
            for fut in concurrent.futures.as_completed(futures):
                idx, text = fut.result()
                page_texts[idx] = text
                if not text.strip():
                    failed_pages.append(idx)

        # Serial retry for pages that came back empty (likely buffer errors)
        if failed_pages:
            logger.info(
                f"[PDF] Retrying {len(failed_pages)} empty pages serially: {failed_pages[:10]}"
            )
            for idx in failed_pages:
                _, text = _extract_page_text((idx, data))
                if text.strip():
                    page_texts[idx] = text

        full_text = "\n".join(page_texts)
        avg_chars_per_page = len(full_text.strip()) / max(n_pages, 1)

        if avg_chars_per_page < 80 and (OCR_AVAILABLE or EASYOCR_AVAILABLE):
            logger.info(
                f"[OCR] PDF has sparse text ({avg_chars_per_page:.0f} chars/page)"
                f" -- running parallel OCR on {n_pages} pages"
            )
            try:
                poppler_kw = {}
                if POPPLER_PATH and os.path.exists(POPPLER_PATH):
                    poppler_kw["poppler_path"] = POPPLER_PATH
                # Lower DPI for very large PDFs to reduce per-image RAM
                dpi = 200 if n_pages > 60 else 300
                images = convert_from_bytes(data, dpi=dpi, **poppler_kw)
                ocr_args: list[tuple[int, bytes]] = []
                for i, img in enumerate(images):
                    buf = io.BytesIO()
                    img.save(buf, format="PNG")
                    ocr_args.append((i, buf.getvalue()))
                ocr_texts = [""] * len(images)
                ocr_workers = max(1, OCR_WORKERS - 1) if n_pages > 40 else OCR_WORKERS
                with concurrent.futures.ThreadPoolExecutor(max_workers=ocr_workers) as exe:
                    futures = {exe.submit(_ocr_page_worker, arg): arg[0] for arg in ocr_args}
                    for fut in concurrent.futures.as_completed(futures):
                        idx, text = fut.result()
                        ocr_texts[idx] = text
                ocr_full = "\n".join(ocr_texts)
                if len(ocr_full.strip()) > len(full_text.strip()):
                    full_text = ocr_full
            except Exception as e:
                logger.error(f"[OCR] Parallel OCR failed: {e}")

        logger.info(f"[PDF] Extracted {len(full_text)} chars from {n_pages} pages")
        return full_text
    except Exception as e:
        logger.error(f"[PDF] Extraction error: {e}")
        return ""


# ─── Other File Types (unchanged) ────────────────────────────────────────────
def extract_text_from_file(file_path: str, file_type: str) -> str:
    try:
        if file_type == "pdf":
            with open(file_path, "rb") as f:
                return _extract_pdf_bytes(f.read())
        elif file_type == "docx":
            doc  = DocxDocument(file_path)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join(c.text.strip() for c in row.cells)
                    if row_text.strip():
                        text += "\n" + row_text
            return text
        elif file_type == "txt":
            with open(file_path, encoding="utf-8", errors="ignore") as f:
                return f.read()
        elif file_type == "xml":
            tree = ET.parse(file_path)
            return " ".join(n.text for n in tree.iter() if n.text)
    except Exception as e:
        logger.error(f"Text extraction failed for {file_path}: {e}")
    return ""


def _ocr_image_url(img_url: str) -> str:
    if not (OCR_AVAILABLE or EASYOCR_AVAILABLE):
        return ""
    resp = _safe_get(img_url)
    if not resp or len(resp.content) / 1024 < MIN_IMAGE_SIZE_KB:
        return ""
    try:
        img  = Image.open(io.BytesIO(resp.content)).convert("RGB")
        return _ocr_single_image(img)
    except Exception as e:
        logger.warning(f"Image OCR error {img_url}: {e}")
        return ""


# ─── Document Type Detection (unchanged) ─────────────────────────────────────
def chunk_general_document(text: str) -> list[dict]:
    CHUNK_SIZE    = 800
    CHUNK_OVERLAP = 100
    chunks, current, count = [], [], 0
    for sentence in re.split(r"(?<=[.!?\n])\s+", text.strip()):
        words = sentence.split()
        if count + len(words) > CHUNK_SIZE and current:
            chunks.append({"text": " ".join(current), "subject_json": None,
                           "semester": None, "content_type": None})
            current = " ".join(current).split()[-CHUNK_OVERLAP:] + words
            count   = len(current)
        else:
            current.extend(words)
            count += len(words)
    if current:
        chunks.append({"text": " ".join(current), "subject_json": None,
                       "semester": None, "content_type": None})
    logger.info(f"General chunking: {len(chunks)} chunks")
    return chunks


def chunk_syllabus_document(text: str, source_url: str = "") -> list[dict]:
    """
    Syllabus-aware chunking for normal PDFs and noisy OCR text.
    Keeps sections like Unit/Module/Outcomes/Scheme together.
    """
    cleaned = (text or "").replace("\r", "\n")
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).strip()
    if len(cleaned) < 80:
        return chunk_general_document(text)

    def _normalize_sem(raw: str) -> str | None:
        if not raw:
            return None
        token = raw.strip().upper()
        sem_map = {
            "1": "SEM-I", "I": "SEM-I",
            "2": "SEM-II", "II": "SEM-II",
            "3": "SEM-III", "III": "SEM-III",
            "4": "SEM-IV", "IV": "SEM-IV",
            "5": "SEM-V", "V": "SEM-V",
            "6": "SEM-VI", "VI": "SEM-VI",
            "7": "SEM-VII", "VII": "SEM-VII",
            "8": "SEM-VIII", "VIII": "SEM-VIII",
        }
        return sem_map.get(token)

    def _clean_course_name(name: str) -> str:
        n = " ".join((name or "").replace("\xe2\x80\x93", "-").replace("\u2013", "-").split())
        n = n.strip(" -:|,")
        n = re.split(r"\s{2,}\d{1,3}(?:\s+\d{1,3}){1,}", n, maxsplit=1)[0].strip()
        tokens = n.split()
        if tokens:
            cut_idx = None
            for i, tok in enumerate(tokens):
                if re.match(r"^(?:\d+(?:\.\d+)?|--+|[-/]+)$", tok):
                    tail = tokens[i:]
                    numeric_like = sum(
                        1 for t in tail
                        if re.match(r"^(?:\d+(?:\.\d+)?|--+|[-/]+|[A-Z]{1,2}\+?)$", t)
                    )
                    if tail and (numeric_like / len(tail)) >= 0.6:
                        cut_idx = i
                        break
            if cut_idx is not None:
                tokens = tokens[:cut_idx]
        n = " ".join(tokens).strip(" -:|,")
        n = re.sub(r"\s{2,}", " ", n)
        return n

    def _is_valid_course_name(name: str) -> bool:
        if not name:
            return False
        if len(name) < 3:
            return False
        if not re.search(r"[A-Za-z]", name):
            return False
        low = name.lower()
        bad = [
            "teaching scheme", "credits assigned", "examination scheme",
            "internal assessment", "course code", "course name", "total",
        ]
        return not any(b in low for b in bad)

    def _extract_semester_courses(src_text: str) -> dict[str, dict[str, str]]:
        courses_by_sem: dict[str, dict[str, str]] = {}
        lines = [re.sub(r"\s+", " ", (ln or "").strip()) for ln in src_text.splitlines()]
        lines = [ln for ln in lines if ln]

        sem_re = re.compile(r"\bSemester\s*(I{1,3}V?|IV|V?I{0,3}|[1-8])\b", re.I)
        code_re = re.compile(r"\b([A-Z]{2,8}\s*\d{3,4}[A-Z]?)\b")

        current_sem = None
        i = 0
        while i < len(lines):
            line = lines[i]

            sem_m = sem_re.search(line)
            if sem_m:
                current_sem = _normalize_sem(sem_m.group(1))
                if current_sem and current_sem not in courses_by_sem:
                    courses_by_sem[current_sem] = {}

            if not current_sem:
                i += 1
                continue

            split_m = re.match(r"^(CSDO|CSDOL|ILO)\b\s*(.+)$", line, re.I)
            if split_m:
                found_split = False
                for j in range(i + 1, min(i + 5, len(lines))):
                    look = lines[j]
                    next_m = re.match(r"^(\d{3}[A-Z]?)\b\s*(.+)$", look, re.I)
                    if not next_m:
                        continue
                    code = f"{split_m.group(1).upper()}{next_m.group(1).upper()}"
                    name = _clean_course_name(f"{split_m.group(2)} {next_m.group(2)}")
                    if _is_valid_course_name(name):
                        courses_by_sem[current_sem][code] = name
                    i = j + 1
                    found_split = True
                    break
                if found_split:
                    continue

            code_m = code_re.search(line)
            if code_m:
                code = re.sub(r"\s+", "", code_m.group(1).upper())
                rem = line[code_m.end():].strip(" -:|,")
                name = _clean_course_name(rem)

                if not _is_valid_course_name(name):
                    prev = lines[i - 1] if i > 0 else ""
                    nxt = lines[i + 1] if i + 1 < len(lines) else ""
                    candidates = [
                        _clean_course_name(f"{prev} {nxt}"),
                        _clean_course_name(prev),
                        _clean_course_name(nxt),
                    ]
                    name = next((c for c in candidates if _is_valid_course_name(c)), "")

                if _is_valid_course_name(name):
                    old = courses_by_sem[current_sem].get(code, "")
                    if not old or len(name) > len(old):
                        courses_by_sem[current_sem][code] = name

            i += 1

        return {k: v for k, v in courses_by_sem.items() if v}

    sem_courses = _extract_semester_courses(cleaned)
    structured_chunks = []
    for sem, courses in sem_courses.items():
        lines = [f"- {code}: {name}" for code, name in sorted(courses.items())]
        summary_text = (
            f"Syllabus Semester: {sem}\n"
            f"Syllabus Courses ({len(lines)}):\n" + "\n".join(lines)
        )
        structured_chunks.append(
            {
                "text": summary_text,
                "subject_json": None,
                "semester": sem,
                "content_type": "SYLLABUS",
                "department": None,
            }
        )
        for code, name in sorted(courses.items()):
            structured_chunks.append(
                {
                    "text": (
                        f"Syllabus Semester: {sem}\n"
                        f"Course Code: {code}\n"
                        f"Course Name: {name}\n"
                        f"Subject Name: {name}"
                    ),
                    "subject_json": None,
                    "semester": sem,
                    "content_type": "SYLLABUS",
                    "department": None,
                }
            )

    heading_re = re.compile(
        r"(?im)^\s*(?:"
        r"UNIT\s*[-:]?\s*(?:\d+|[IVX]+)\b|"
        r"MODULE\s*[-:]?\s*\d+\b|"
        r"COURSE\s+OUTCOMES?\b|"
        r"PROGRAM\s+OUTCOMES?\b|"
        r"TEACHING\s+SCHEME\b|"
        r"EXAMINATION\s+SCHEME\b|"
        r"SYLLABUS\b"
        r")"
    )
    matches = list(heading_re.finditer(cleaned))

    sections = []
    if matches:
        for i, m in enumerate(matches):
            start = m.start()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(cleaned)
            sec = cleaned[start:end].strip()
            if len(sec) >= 60:
                sections.append(sec)
    if not sections:
        sections = [cleaned]

    sem = classify_chunk(cleaned[:2000], source_url=source_url).get("semester")
    chunks = []
    for sec in sections:
        words = sec.split()
        if len(words) > 260:
            step = 220
            overlap = 30
            idx = 0
            while idx < len(words):
                part = " ".join(words[idx:idx + step]).strip()
                if part:
                    chunks.append({
                        "text": part,
                        "subject_json": None,
                        "semester": sem,
                        "content_type": "SYLLABUS",
                        "department": None,
                    })
                if idx + step >= len(words):
                    break
                idx += max(1, step - overlap)
        else:
            chunks.append({
                "text": sec,
                "subject_json": None,
                "semester": sem,
                "content_type": "SYLLABUS",
                "department": None,
            })

    if structured_chunks:
        chunks = structured_chunks + chunks

    logger.info(
        f"Syllabus chunking: {len(chunks)} chunks | sem_structured={len(structured_chunks)} | source={source_url[:80]}"
    )
    return chunks


# ─── PDF type detection (unchanged) ──────────────────────────────────────────
PDF_SINGLE_CHUNK_KEYWORDS = [
    "question paper","question-paper","qp",
    "topper","topper-list","toppers",
    "newsletter","news-letter","news_letter",
    "activity","activities",
    "project","internship",
    "circular","notice",
    "placement","recruit",
]

def _is_single_chunk_pdf(source_url: str) -> bool:
    url_lower = (source_url or "").lower()
    return any(kw in url_lower for kw in PDF_SINGLE_CHUNK_KEYWORDS)


def _build_chunks_for_classified_document(
    text: str,
    source_url: str = "",
    pdf_data: bytes = None,
    content_type: str = None,
) -> list[dict]:
    ctype = (content_type or classify_chunk(text[:2000], source_url=source_url).get("content_type") or "GENERAL").upper()

    if ctype == "RESULT":
        semester = detect_semester_from_text(text)
        students = parse_result_document(
            text=text,
            pdf_data=pdf_data,
            semester=semester,
            min_students=1,
        )
        if students:
            chunks = students_to_result_chunks(students)
            logger.info(f"Result extractor triggered | chunks={len(chunks)} | source={source_url[:80]}")
            return chunks
        logger.warning(f"Result content detected but no student rows extracted | source={source_url[:80]}")

    if ctype == "SYLLABUS":
        return chunk_syllabus_document(text, source_url=source_url)

    if source_url and _is_single_chunk_pdf(source_url):
        logger.info(f"Single-chunk PDF: {source_url[:60]}")
        preview         = text[:1500].strip()
        filename        = source_url.split("/")[-1]
        chunk_text_val  = f"[PDF Document] {filename}\nURL: {source_url}\n\n{preview}"
        return [{
            "text":         chunk_text_val,
            "subject_json": json.dumps({"url": source_url, "type": "pdf"}),
            "semester":     None,
            "content_type": None,
            "department":   None,
            "source_url":   source_url,
        }]

    logger.info(f"General chunking path | type={ctype} | source={source_url[:80]}")
    return chunk_general_document(text)


def chunk_text(text: str, pdf_data: bytes = None, source_url: str = "") -> list[dict]:
    return _build_chunks_for_classified_document(
        text=text,
        source_url=source_url,
        pdf_data=pdf_data,
        content_type=None,
    )


# ─── Embeddings + FAISS (unchanged) ──────────────────────────────────────────
def create_embeddings(texts: list[str]) -> np.ndarray:
    if not texts:
        return np.array([])
    try:
        from app.chat.cache import get_cached_embedding, set_cached_embedding
        results = []
        miss_texts = []
        miss_idx   = []
        for i, t in enumerate(texts):
            cached = get_cached_embedding(t)
            if cached is not None:
                results.append((i, cached))
            else:
                miss_texts.append(t)
                miss_idx.append(i)
        if miss_texts:
            model = get_model()
            new_embs = model.encode(miss_texts, normalize_embeddings=True,
                                    show_progress_bar=False).astype("float32")
            for j, (orig_i, text) in enumerate(zip(miss_idx, miss_texts)):
                emb_row = new_embs[j:j+1]
                set_cached_embedding(text, emb_row)
                results.append((orig_i, emb_row))
        results.sort(key=lambda x: x[0])
        emb = np.vstack([r[1] for r in results]).astype("float32")
    except ImportError:
        model = get_model()
        emb   = model.encode(texts, normalize_embeddings=True, show_progress_bar=False)
        emb   = np.array(emb).astype("float32")
    return emb.reshape(1, -1) if emb.ndim == 1 else emb


def save_to_faiss(embeddings: np.ndarray, chunk_ids: list[int]):
    import faiss
    if embeddings is None or len(embeddings) == 0:
        return
    index = get_index()
    emb   = np.array(embeddings).astype("float32")
    if emb.ndim == 1:
        emb = emb.reshape(1, -1)
    if emb.shape[1] != index.d:
        logger.error("Embedding dimension mismatch")
        return
    index.add_with_ids(emb, np.array(chunk_ids, dtype=np.int64))
    save_index(index)


# ─── Core Save (unchanged) ───────────────────────────────────────────────────
def _save_chunks(chunks: list[dict], document_id: int, db,
                 source_url: str = "") -> int:
    if not chunks:
        return 0
    meta       = classify_chunk(" ".join(c["text"] for c in chunks[:5]), source_url)
    embeddings = create_embeddings([c["text"] for c in chunks])
    chunk_ids  = []
    for i, c in enumerate(chunks):
        db_chunk = DocumentChunk(
            document_id  = document_id,
            chunk_text   = c["text"],
            chunk_index  = i,
            vector_id    = None,
            department   = c.get("department")   or meta["department"],
            content_type = c.get("content_type") or meta["content_type"],
            semester     = c.get("semester")     or meta["semester"],
            subject_data = c.get("subject_json"),
            source_url   = source_url or None,
        )
        db.add(db_chunk)
        db.flush()
        chunk_ids.append(db_chunk.id)
    db.commit()
    for cid in chunk_ids:
        ch = db.query(DocumentChunk).get(cid)
        if ch:
            ch.vector_id = cid
    db.commit()
    save_to_faiss(embeddings, chunk_ids)
    try:
        from app.chat.hybrid_search import invalidate_bm25
        invalidate_bm25()
    except Exception:
        pass
    logger.info(f"Saved {len(chunks)} chunks | {meta['content_type']} | {source_url[:60]}")
    return len(chunks)


# ─── Process Local File (unchanged) ──────────────────────────────────────────
def process_document(file_path: str, file_type: str, document_id: int) -> dict:
    db = SessionLocal()
    try:
        pdf_data = None
        if file_type == "pdf":
            with open(file_path, "rb") as f:
                pdf_data = f.read()
            text = _extract_pdf_bytes(pdf_data)
        else:
            text = extract_text_from_file(file_path, file_type)
        if not text or len(text.strip()) < 50:
            return {"status": "error", "message": "Insufficient text extracted"}
        meta = classify_chunk(text[:2000], source_url=file_path)
        chunks = _build_chunks_for_classified_document(
            text=text[:MAX_DOC_CHARS],
            source_url=file_path,
            pdf_data=pdf_data,
            content_type=meta.get("content_type"),
        )
        if not chunks:
            return {"status": "error", "message": "No chunks generated"}
        total = _save_chunks(chunks, document_id, db, source_url=file_path)
        try:
            from app.chat.cache import invalidate_response_cache
            invalidate_response_cache()
        except Exception:
            pass
        return {"status": "success", "chunks_processed": total}
    except Exception as e:
        db.rollback()
        logger.error(f"process_document error: {e}")
        return {"status": "error", "message": str(e)}
    finally:
        db.close()


# ─── Smart Result Reindex (unchanged) ────────────────────────────────────────
def reindex_result_document(document_id: int) -> dict:
    db = SessionLocal()
    try:
        from app.models.document import Document
        doc = db.query(Document).filter(Document.id == document_id).first()
        if not doc:
            return {"status": "error", "message": f"Document {document_id} not found"}
        if not doc.file_path or not os.path.exists(doc.file_path):
            return {"status": "error", "message": f"File not found: {doc.file_path}"}
        with open(doc.file_path, "rb") as f:
            pdf_data = f.read()
        text = _extract_pdf_bytes(pdf_data)
        meta = classify_chunk(text[:2000], source_url=doc.source_url or doc.filename or "")
        if (meta.get("content_type") or "").upper() != "RESULT":
            return {"status": "error", "message": "Not a result PDF"}
        semester = meta.get("semester") or detect_semester_from_text(text)
        students = parse_result_document(
            text=text,
            pdf_data=pdf_data,
            semester=semester,
            min_students=5,
        )
        if not students:
            return {"status": "error", "message": "No students parsed"}
        db.query(DocumentChunk).filter(DocumentChunk.document_id == document_id).delete()
        db.commit()
        chunks = students_to_result_chunks(students)
        new_chunk_ids = []
        for i, c in enumerate(chunks):
            dc = DocumentChunk(
                document_id=document_id, chunk_text=c["text"],
                chunk_index=i, vector_id=None,
                department=c.get("department","DS"),
                content_type="RESULT",
                semester=c.get("semester") or semester,
                subject_data=c.get("subject_json"),
            )
            db.add(dc)
            db.flush()
            new_chunk_ids.append(dc.id)
        db.commit()
        for cid in new_chunk_ids:
            ch = db.query(DocumentChunk).get(cid)
            if ch: ch.vector_id = cid
        db.commit()
        import faiss as fl
        all_chunks = db.query(DocumentChunk).all()
        emb = create_embeddings([c.chunk_text for c in all_chunks])
        ids = np.array([c.id for c in all_chunks], dtype="int64")
        idx = fl.IndexIDMap(fl.IndexFlatIP(emb.shape[1]))
        idx.add_with_ids(emb, ids)
        save_index(idx)
        try:
            from app.chat.hybrid_search import invalidate_bm25
            invalidate_bm25()
        except Exception:
            pass
        logger.info(f"Reindexed result doc {document_id}: {len(students)} students")
        return {"status": "success", "students_found": len(students),
                "chunks_created": len(new_chunk_ids), "semester": semester}
    except Exception as e:
        db.rollback()
        logger.error(f"reindex error: {e}")
        return {"status": "error", "message": str(e)}
    finally:
        db.close()


# ─── Web Scraper Helpers (unchanged) ─────────────────────────────────────────
def _get_extension(url: str) -> str:
    _, ext = os.path.splitext(urlparse(url).path)
    return ext.lstrip(".").lower()


def _url_hash(url: str) -> str:
    return hashlib.md5(url.encode()).hexdigest()


def _should_skip_url(url: str) -> bool:
    if any(re.search(p, url, re.IGNORECASE) for p in SKIP_URL_PATTERNS):
        return True
    blocked = _get_blocked_domains()
    if blocked:
        url_lower = url.lower()
        if any(b in url_lower for b in blocked):
            logger.debug(f"[BLOCKED] Skipping: {url[:80]}")
            return True
    return False


def _safe_get(url: str, stream: bool = False):
    if not url.startswith("http"):
        return None
    for attempt in range(3):
        try:
            resp = requests.get(url, headers=HEADERS,
                                timeout=REQUEST_TIMEOUT, stream=stream)
            resp.raise_for_status()
            return resp
        except requests.RequestException as e:
            logger.warning(f"Request attempt {attempt+1} failed {url}: {e}")
            time.sleep(1.5 ** attempt)
    return None


def _extract_links(soup: BeautifulSoup, base_url: str,
                   base_domain: str) -> list[str]:
    links = []
    for tag in soup.find_all("a", href=True):
        href = tag["href"].strip()
        if not href or href.startswith(("javascript:", "mailto:", "tel:", "#")):
            continue
        full = urljoin(base_url, href).split("#")[0].rstrip("/")
        if not full.startswith("http"):
            continue
        if urlparse(full).netloc != base_domain:
            continue
        if _should_skip_url(full):
            continue
        links.append(full)
    return links


# ─── Web Scraper helpers (unchanged) ─────────────────────────────────────────
def _save_pdf_to_disk(pdf_bytes: bytes, source_url: str) -> str:
    pdf_dir = os.path.join(os.getenv("UPLOAD_DIR", "data"), "pdfs")
    os.makedirs(pdf_dir, exist_ok=True)
    orig_name = source_url.split("/")[-1].split("?")[0]
    orig_name = re.sub(r"[^a-zA-Z0-9._-]", "_", orig_name)
    if not orig_name.lower().endswith(".pdf"):
        orig_name += ".pdf"
    file_path = os.path.join(pdf_dir, orig_name)
    counter = 1
    base, ext = os.path.splitext(file_path)
    while os.path.exists(file_path):
        file_path = f"{base}_{counter}{ext}"
        counter += 1
    try:
        with open(file_path, "wb") as f:
            f.write(pdf_bytes)
        logger.info(f"[PDF] Saved to disk: {file_path} ({len(pdf_bytes)//1024}KB)")
        return file_path
    except Exception as e:
        logger.error(f"[PDF] Failed to save {orig_name}: {e}")
        return ""


def _register_pdf_document(db, source_url: str, file_path: str,
                           pdf_bytes: bytes, text: str, parent_doc_id: int) -> int:
    from app.models.document import Document
    from app.document.auto_label import auto_label
    file_hash = hashlib.sha256(pdf_bytes).hexdigest()
    filename = os.path.basename(file_path)
    label_info = auto_label(filename, source_url=source_url)
    existing = db.query(Document).filter(Document.file_hash == file_hash).first()
    if existing:
        if not existing.file_path or not os.path.exists(existing.file_path):
            existing.file_path = file_path
            db.commit()
        return existing.id
    meta = classify_chunk(text[:1000], source_url=source_url)
    doc = Document(
        filename=filename, file_path=file_path, file_type="pdf",
        file_hash=file_hash,
        department=label_info.get("dept_tag") or meta["department"] or "GENERAL",
        semester=0, subject="GENERAL", uploaded_by=1,
        source_url=source_url,
        source_label=label_info.get("source_label") or filename.replace(".pdf","").replace("_"," ").strip(),
        dept_tag=label_info.get("dept_tag") or meta["department"],
        is_active=True,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    logger.info(f"[PDF] Registered Document id={doc.id} | {filename}")
    return doc.id


# ─── helpers shared by the crawl loop ────────────────────────────────────────

def _collect_page_links(
    url: str,
    soup: "BeautifulSoup",
    visited: set,
    base_domain: str,
) -> tuple[list[str], list[str], list[str]]:
    """
    Walk every <a> on the page and bucket links into three lists:
      pdf_doc_links  — PDF / docx / txt / xml  (download & index immediately)
      image_links    — image files              (OCR if available)
      html_links     — HTML pages               (ask for confirmation)
    Already-visited or skip-pattern URLs are excluded.
    """
    pdf_doc_links: list[str] = []
    image_links:  list[str] = []
    html_links:   list[str] = []

    for a_tag in soup.find_all("a", href=True):
        href = a_tag["href"].strip()
        if not href:
            continue
        full = urljoin(url, href).split("#")[0].rstrip("/")
        if not full.startswith("http"):
            continue
        if full in visited:
            continue
        if _should_skip_url(full):
            continue
        if urlparse(full).netloc != base_domain:
            continue

        ext = _get_extension(full)
        if ext == "pdf" or ext in {"docx", "txt", "xml"}:
            pdf_doc_links.append(full)
        elif ext in {"jpg", "jpeg", "png", "gif", "webp", "bmp"}:
            image_links.append(full)
        else:
            html_links.append(full)

    return pdf_doc_links, image_links, html_links


def _scrape_html_page(
    url: str,
    resp,
    visited: set,
    seen_hashes: set,
    base_domain: str,
    summary: dict,
    db,
    get_doc_id,
    save_pdf_fn,
    save_fn,
) -> list[str]:
    """
    Fully process one HTML page in order:

      Step 1 — Extract the full page text and save it immediately.
      Step 2 — Download and fully extract the first 2 PDF/doc links found on
               the page, then download and index all remaining PDF/doc links.
               Every document is processed completely before moving on.
      Step 3 — OCR inline <img> tags (if OCR engine available).
      Step 4 — Return new HTML page URLs found on this page so the caller
               can ask the user for confirmation before enqueuing them.
    """
    try:
        soup = BeautifulSoup(resp.text, "lxml")
    except Exception:
        soup = BeautifulSoup(resp.text, "html.parser")

    # ── strip boilerplate ────────────────────────────────────────────────────
    for tag in soup(
        ["script", "style", "nav", "footer", "header",
         "aside", "noscript", "iframe", "form"]
    ):
        tag.decompose()
    for sel in [".cookie-banner", ".popup", "#cookie"]:
        for el in soup.select(sel):
            el.decompose()

    # ── Step 1: extract & save full page text ────────────────────────────────
    main = (
        soup.find("main")
        or soup.find("article")
        or soup.find(id="content")
        or soup.find(class_="content")
        or soup.find(id="main")
        or soup.body
    )
    page_text = (main or soup).get_text(separator="\n", strip=True)
    page_text = re.sub(r"\n{3,}", "\n\n", page_text)
    if len(page_text) > MAX_HTML_CHARS:
        page_text = page_text[:MAX_HTML_CHARS]

    save_fn(url, page_text, "html")
    logger.info(f"[SCRAPE] Saved page text ({len(page_text)} chars) | {url[:80]}")

    # ── collect all links bucketed by type ───────────────────────────────────
    pdf_doc_links, image_links, html_links = _collect_page_links(
        url, soup, visited, base_domain
    )
    logger.info(
        f"[SCRAPE] Found on page: {len(pdf_doc_links)} docs, "
        f"{len(image_links)} images, {len(html_links)} HTML links | {url[:60]}"
    )

    # ── Step 2: process PDF/doc links — first 2 are fully extracted first ────
    priority   = pdf_doc_links[:2]
    remaining  = pdf_doc_links[2:]

    def _process_doc_link(doc_url: str):
        if doc_url in visited:
            return
        if _should_skip_url(doc_url):
            logger.debug(f"[BLOCKED] Skipping doc: {doc_url[:80]}")
            return
        visited.add(doc_url)
        ext = _get_extension(doc_url)
        time.sleep(REQUEST_DELAY)

        if ext == "pdf":
            logger.info(f"[PDF] Downloading: {doc_url[:80]}")
            pdf_resp = _safe_get(doc_url)
            if pdf_resp and pdf_resp.content:
                save_pdf_fn(doc_url, pdf_resp.content)

        elif ext in {"docx", "txt", "xml"}:
            logger.info(f"[DOC] Downloading: {doc_url[:80]}")
            doc_resp = _safe_get(doc_url)
            if doc_resp:
                tmp = os.path.join(
                    os.environ.get("TEMP", "/tmp"),
                    _url_hash(doc_url) + "." + ext,
                )
                with open(tmp, "wb") as f:
                    f.write(doc_resp.content)
                doc_text = extract_text_from_file(tmp, ext)
                try:
                    os.remove(tmp)
                except Exception:
                    pass
                if doc_text:
                    save_fn(doc_url, doc_text[:MAX_DOC_CHARS], ext)

    # Priority: first 2 docs fully extracted before anything else
    for doc_url in priority:
        logger.info(f"[PRIORITY-DOC] Processing first-2: {doc_url[:80]}")
        _process_doc_link(doc_url)

    # Remaining docs on the page
    for doc_url in remaining:
        _process_doc_link(doc_url)

    # ── Step 3: OCR inline images ────────────────────────────────────────────
    if OCR_AVAILABLE or EASYOCR_AVAILABLE:
        for img_tag in soup.find_all("img", src=True):
            img_url = urljoin(url, img_tag["src"]).rstrip("/")
            if img_url not in visited and not _should_skip_url(img_url):
                visited.add(img_url)
                img_text = _ocr_image_url(img_url)
                if img_text:
                    save_fn(img_url, img_text, "image")

    # OCR image <a> links
    for img_url in image_links:
        if img_url not in visited:
            visited.add(img_url)
            img_text = _ocr_image_url(img_url)
            if img_text:
                save_fn(img_url, img_text, "image")

    # ── Step 4: return HTML links for caller to confirm ──────────────────────
    return html_links


# ─── Web Scraper ─────────────────────────────────────────────────────────────


def process_website(start_url: str, document_id_map: dict = None) -> dict:
    """
    Crawl strategy
    ──────────────
    For every HTML page dequeued:

      1. Extract the full page text and save it immediately.
      2. Download + fully extract the FIRST 2 PDF/doc links on the page,
         then download + index all remaining PDF/doc links on the page.
         Every document is drained completely before the crawler moves on.
      3. OCR inline images on the page.
      4. Child HTML links are AUTO-FILTERED:
           - URL matches a DS/CSE/AIML keyword  →  enqueued automatically
           - URL does not match                 →  silently skipped
         No manual confirmation needed.
    """
    db          = SessionLocal()
    visited     = set()
    seen_hashes = set()
    queue: deque = deque([(start_url.rstrip("/"), 0)])
    base_domain  = urlparse(start_url).netloc

    summary = {
        "html": 0, "pdf": 0, "docx": 0, "image": 0,
        "RESULT": 0, "FACULTY": 0, "NOTICE": 0, "SYLLABUS": 0,
        "FEE": 0, "PLACEMENT": 0, "GENERAL": 0, "errors": 0,
        "skipped_duplicate": 0, "pdfs_saved": 0,
        "skipped_irrelevant": 0,
    }

    # ── helpers ───────────────────────────────────────────────────────────────
    def get_doc_id(file_type: str) -> int:
        if not document_id_map:
            return 0
        return document_id_map.get(file_type, document_id_map.get("default", 0))

    def save_pdf(source_url: str, pdf_bytes: bytes):
        """Download, register and index a single PDF."""
        if _should_skip_url(source_url):
            logger.debug(f"[BLOCKED] Skipping PDF: {source_url[:80]}")
            return
        if not pdf_bytes or len(pdf_bytes) < 1024:
            return
        content_hash = hashlib.md5(pdf_bytes[:8192]).hexdigest()
        if content_hash in seen_hashes:
            summary["skipped_duplicate"] += 1
            return
        seen_hashes.add(content_hash)

        text = _extract_pdf_bytes(pdf_bytes)
        if len(text.strip()) < 30:
            text = f"[PDF] {source_url.split('/')[-1]}"

        local_path = _save_pdf_to_disk(pdf_bytes, source_url)
        if not local_path:
            return

        parent_id = get_doc_id("pdf")
        try:
            doc_id = _register_pdf_document(
                db, source_url, local_path, pdf_bytes, text, parent_id
            )
        except Exception as e:
            logger.warning(f"[PDF] Register failed {source_url}: {e}")
            db.rollback()
            return

        try:
            meta = classify_chunk(text[:500], source_url=source_url)
            if len(text.strip()) >= 50:
                chunks = _build_chunks_for_classified_document(
                    text=text[:MAX_DOC_CHARS],
                    source_url=source_url,
                    pdf_data=pdf_bytes,
                    content_type=meta.get("content_type"),
                )
            else:
                filename  = os.path.basename(local_path)
                stub_text = (
                    f"[PDF Document] {filename}\n"
                    f"URL: {source_url}\nLocal: available for download"
                )
                chunks = [{
                    "text": stub_text, "subject_json": None,
                    "semester": None, "content_type": None,
                }]
            _save_chunks(chunks, doc_id, db, source_url=source_url)
            ctype = meta["content_type"]
            summary[ctype]         = summary.get(ctype, 0) + len(chunks)
            summary["pdf"]        += 1
            summary["pdfs_saved"] += 1
        except Exception as e:
            summary["errors"] += 1
            logger.warning(f"[PDF] Index error {source_url}: {e}")
            db.rollback()

    def save(source_url: str, text: str, file_type: str = "html",
             pdf_data: bytes = None):
        """Index any non-PDF text content."""
        if not text or len(text.strip()) < 100:
            return
        content_hash = hashlib.md5(text[:5000].encode()).hexdigest()
        if content_hash in seen_hashes:
            summary["skipped_duplicate"] += 1
            return
        seen_hashes.add(content_hash)

        doc_id = get_doc_id(file_type)
        try:
            meta   = classify_chunk(text[:500], source_url=source_url)
            chunks = _build_chunks_for_classified_document(
                text=text,
                source_url=source_url,
                pdf_data=pdf_data,
                content_type=meta.get("content_type"),
            )
            n     = _save_chunks(chunks, doc_id, db, source_url=source_url)
            ctype = meta["content_type"]
            summary[ctype]     = summary.get(ctype, 0) + n
            summary[file_type] = summary.get(file_type, 0) + 1
        except Exception as e:
            summary["errors"] += 1
            logger.warning(f"Save error {source_url}: {e}")
            db.rollback()

    # ── main crawl loop ───────────────────────────────────────────────────────
    pages_crawled = 0
    try:
        while queue and pages_crawled < MAX_PAGES_PER_CRAWL:
            url, depth = queue.popleft()
            url = url.rstrip("/")

            if url in visited or _should_skip_url(url):
                continue
            visited.add(url)

            ext = _get_extension(url)
            logger.info(f"[CRAWL] Page {pages_crawled + 1} | depth={depth} | {url[:80]}")

            # ── direct file URLs ──────────────────────────────────────────────
            if ext == "pdf":
                pdf_resp = _safe_get(url)
                if pdf_resp and pdf_resp.content:
                    save_pdf(url, pdf_resp.content)
                pages_crawled += 1
                time.sleep(REQUEST_DELAY)
                continue

            if ext in {"docx", "txt", "xml"}:
                doc_resp = _safe_get(url)
                if doc_resp:
                    tmp = os.path.join(
                        os.environ.get("TEMP", "/tmp"),
                        _url_hash(url) + "." + ext,
                    )
                    with open(tmp, "wb") as f:
                        f.write(doc_resp.content)
                    text = extract_text_from_file(tmp, ext)
                    try:
                        os.remove(tmp)
                    except Exception:
                        pass
                    if text:
                        save(url, text[:MAX_DOC_CHARS], ext)
                pages_crawled += 1
                time.sleep(REQUEST_DELAY)
                continue

            if ext in {"jpg", "jpeg", "png", "gif", "webp", "bmp"}:
                img_text = _ocr_image_url(url)
                if img_text:
                    save(url, img_text, "image")
                pages_crawled += 1
                time.sleep(REQUEST_DELAY)
                continue

            # ── fetch HTML page ───────────────────────────────────────────────
            resp = _safe_get(url)
            if not resp:
                pages_crawled += 1
                continue

            ct = resp.headers.get("Content-Type", "")

            if "application/pdf" in ct:
                save_pdf(url, resp.content)
                pages_crawled += 1
                time.sleep(REQUEST_DELAY)
                continue

            if "text/html" not in ct:
                pages_crawled += 1
                continue

            # ── FULLY process this HTML page (text + all docs + images) ──────
            new_html_links = _scrape_html_page(
                url=url,
                resp=resp,
                visited=visited,
                seen_hashes=seen_hashes,
                base_domain=base_domain,
                summary=summary,
                db=db,
                get_doc_id=get_doc_id,
                save_pdf_fn=save_pdf,
                save_fn=save,
            )
            pages_crawled += 1
            time.sleep(REQUEST_DELAY)

            # ── Step 4: enqueue all child HTML links ──────────────────────────
            if new_html_links and depth < MAX_CRAWL_DEPTH:
                added = 0
                for link in new_html_links:
                    if link not in visited:
                        queue.append((link, depth + 1))
                        added += 1
                logger.info(
                    f"[CRAWL] {added} child links enqueued | parent: {url[:60]}"
                )

    finally:
        # ── rebuild FAISS index ───────────────────────────────────────────────
        try:
            import faiss as fl
            import numpy as _np
            all_chunks = db.query(DocumentChunk).all()
            if all_chunks:
                emb = create_embeddings([c.chunk_text for c in all_chunks])
                ids = _np.array([c.id for c in all_chunks], dtype="int64")
                idx = fl.IndexIDMap(fl.IndexFlatIP(emb.shape[1]))
                idx.add_with_ids(emb, ids)
                save_index(idx)
                logger.info(f"FAISS rebuilt after scrape | vectors={idx.ntotal}")
        except Exception as e:
            logger.error(f"FAISS rebuild error: {e}")
        finally:
            db.close()

    try:
        from app.chat.cache import invalidate_response_cache
        from app.chat.hybrid_search import invalidate_bm25
        invalidate_response_cache()
        invalidate_bm25()
    except Exception:
        pass

    logger.info(f"[CRAWL] Done | pages={pages_crawled} | summary={summary}")
    return {"summary": summary}